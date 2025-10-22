import streamlit as st
import requests
import json
from supabase import create_client, Client
from datetime import datetime, timezone, timedelta
import base64
import docx
import PyPDF2
from io import BytesIO
import diff_match_patch as dmp_module
from docx.shared import RGBColor

# --- .envからキーを読み込む手動関数 ---
def load_env_variables():
    env_vars = {}
    try:
        # pagesフォルダから実行されることを想定し、一つ上の階層の.envを指す
        with open('.env', 'r', encoding='utf-8') as f:
            for line in f:
                if '=' in line:
                    key, value = line.strip().split('=', 1)
                    env_vars[key.strip()] = value.strip().strip('"')
        return env_vars
    except FileNotFoundError:
        return None

# --- Secrets/環境変数からのキー読み込み ---
api_key = None
supabase_url = None
supabase_key = None

try:
    api_key = st.secrets["GEMINI_API_KEY"]
    supabase_url = st.secrets["SUPABASE_URL"]
    supabase_key = st.secrets["SUPABASE_KEY"]
except Exception:
    env_vars = load_env_variables()
    if env_vars:
        api_key = env_vars.get("GEMINI_API_KEY")
        supabase_url = env_vars.get("SUPABASE_URL")
        supabase_key = env_vars.get("SUPABASE_KEY")

if not api_key:
    st.error("GEMINI_API_KEYを読み込めませんでした。Secretsまたは.envファイルを確認してください。")
    st.stop()

# --- ページコンテンツ ---
st.header("📖 就業規則の改正")

with st.form("rule_amendment_form"):
    st.subheader("改正対象のファイルと、改正のテーマを入力してください")
    rule_file = st.file_uploader("改正対象の就業規則ファイル", type=['docx', 'pdf', 'txt', 'md'])
    law_amendment_points = st.text_area("改正のテーマ（キーワード）", height=150, placeholder="（例）育児・介護休業法の改正に対応したい。特に、子の看護休暇の時間単位取得について。...")
    reference_file = st.file_uploader("参考資料（法改正の条文や解説記事など）", type=['txt', 'md'])
    rule_submitted = st.form_submit_button("分析を開始する")

if rule_submitted:
    if not rule_file or not law_amendment_points:
        st.error("改正対象のファイルと、改正のテーマの両方を入力してください。")
    else:
        with st.spinner("改正案とリスクを分析しています..."):
            current_rule = ""
            try:
                if "word" in rule_file.type or "document" in rule_file.type:
                    try:
                        document = docx.Document(BytesIO(rule_file.getvalue()))
                        for para in document.paragraphs:
                            current_rule += para.text + "\n"
                    except Exception:
                        st.error("Wordファイルの読み込みに失敗しました。ファイルが破損しているか、有効なWord形式ではない可能性があります。")
                        st.stop()
                elif "pdf" in rule_file.type:
                    pdf_reader = PyPDF2.PdfReader(BytesIO(rule_file.getvalue()))
                    for page in pdf_reader.pages:
                        current_rule += page.extract_text() + "\n"
                else: # txt, md
                    current_rule = rule_file.getvalue().decode('utf-8')
            except Exception as e:
                st.error(f"就業規則ファイルの読み込みに失敗しました: {e}")
                st.stop()

            reference_text = ""
            if reference_file is not None:
                try:
                    reference_text = reference_file.getvalue().decode('utf-8')
                    st.info("参考資料を読み込み、AIへの指示に含めました。")
                except Exception as e:
                    st.warning(f"参考資料の読み込みに失敗しました: {e}")

            rule_prompt = f'''あなたは、日本の労働法規に精通し、高度なWeb検索能力を持つ、非常に優秀な社会保険労務士AIです。

# あなたのタスク
あなたは、就業規則の文章を改訂する専門家です。以下の指示に従って、正確にタスクを実行してください。

1.  以下の「現行の就業規則の全文」と「改正のテーマ」を理解します。
2.  「改正のテーマ」に基づき、関連する最新の法律、法令、判例などをWebで検索し、どのような修正が必要かを判断します。
3.  **重要：** 「現行の就業規則の全文」の構造（章、条、項、インデント、改行など）を**完全に維持**したまま、「改正のテーマ」とWeb検索結果を反映させるために、**必要な箇所だけを修正**してください。改行の位置も、元の文章に可能な限り忠実に合わせてください。
4.  **変更が不要な条文は、一字一句変えずに、そのまま出力に含めてください。**
5.  最終的に、**修正後の就業規則の全文**を `new_rule_draft` として返してください。
6.  また、その改正によって生じる可能性のある、企業側の法的リスクや実務上の課題も指摘してください。

# 入力情報
## 現行の就業規則の全文
```text
{current_rule}
```

## 改正のテーマ（キーワード）
```text
{law_amendment_points}
```

## （任意）ユーザー提供の参考資料
```text
{reference_text}
```

# 出力形式
必ず以下のJSON形式で出力してください。
```json
{{
  "new_rule_draft": "ここに、元の構造を完全に維持した、修正後の就業規則の全文を記述",
  "risks_and_issues": [
    {{
      "type": "法的リスク or 実務上の課題",
      "description": "具体的なリスクや課題内容を記述"
    }}
  ]
}}
```'''
            rule_payload = {
                "contents": [{"parts": [{"text": rule_prompt}]}] ,
                "generationConfig": {"temperature": 0.7, "maxOutputTokens": 8192},
                "safetySettings": [{"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}, {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}, {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}, {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"}]
            }

            try:
                url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent?key={api_key}"
                response = requests.post(url, json=rule_payload, timeout=180)
                response.raise_for_status()
                result_json = response.json()
                raw_text = result_json['candidates'][0]['content']['parts'][0]['text']

                try:
                    start_index = raw_text.index('{')
                    end_index = raw_text.rindex('}')
                    json_text = raw_text[start_index:end_index+1]
                    data = json.loads(json_text)
                    st.success("分析が完了しました。")
                    st.subheader("新しい条文案")
                    st.markdown(data["new_rule_draft"])
                    st.subheader("リスクと確認事項")
                    st.table(data["risks_and_issues"])

                    # --- 変更箇所をハイライトしたWordファイル出力機能 ---
                    st.subheader("変更履歴付きWordファイルとしてダウンロード")
                    try:
                        dmp = dmp_module.diff_match_patch()
                        diffs = dmp.diff_main(current_rule, data["new_rule_draft"])
                        dmp.diff_cleanupSemantic(diffs)

                        document = docx.Document()
                        document.add_heading("就業規則 改正案（変更履歴付き）", 0)
                        p = document.add_paragraph()

                        for op, text in diffs:
                            cleaned_text = text.strip()
                            cleaned_text = cleaned_text.replace(' \n', '\n')
                            run = p.add_run(cleaned_text)
                            if op == dmp.DIFF_DELETE:
                                run.font.strike = True
                                run.font.color.rgb = RGBColor(255, 0, 0)
                            elif op == dmp.DIFF_INSERT:
                                run.font.underline = True
                                run.font.color.rgb = RGBColor(0, 0, 255)

                        bio = BytesIO()
                        document.save(bio)

                        st.download_button(
                            label="改正案（変更履歴付き）をWordでダウンロード",
                            data=bio.getvalue(),
                            file_name="shuugyou_kisoku_kaiseian_diff.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                    except Exception as e:
                        st.error(f"Wordファイルの生成中にエラーが発生しました: {e}")
                        st.exception(e)

                except (ValueError, json.JSONDecodeError):
                    st.error("応答からJSONを抽出できませんでした。")
                    st.text(raw_text)

            except requests.exceptions.HTTPError as http_err:
                if http_err.response.status_code == 503:
                    st.warning("一時的な問題が発生しました。Google Gemini APIが現在利用できません。数分後に再試行してください。")
                else:
                    st.error(f"AIによる分析中にHTTPエラーが発生しました: {http_err}")
                    st.exception(http_err)
            except Exception as e:
                st.error(f"AIによる分析中にエラーが発生しました: {e}")
                st.exception(e)
